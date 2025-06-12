import os
import re
import ssl
import pandas as pd
from docx import Document
from transformers import pipeline

# Ignorar verificación de certificados
ssl._create_default_https_context = ssl._create_unverified_context

# Cargar modelo de emociones
emotion_pipeline = pipeline("text-classification", model="j-hartmann/emotion-english-distilroberta-base", top_k=1)

# Rutas
ruta_entrada = r"C:\DATA\PERSONAL\UNIVERSIDAD\5 NIVEL\HERRAMIENTAS PARA CIENCIA DE DATOS\PROYECTO\TEXTO"
ruta_salida = r"C:\DATA\PERSONAL\UNIVERSIDAD\5 NIVEL\HERRAMIENTAS PARA CIENCIA DE DATOS\PROYECTO\RESULTADOS"
archivo_docx = os.path.join(ruta_salida, "Limpieza_emociones.docx")
archivo_excel = os.path.join(ruta_salida, "Limpieza_emociones.xlsx")

# Leer archivo Word
def leer_docx(ruta_archivo):
    doc = Document(ruta_archivo)
    texto = ""
    for p in doc.paragraphs:
        texto += p.text.strip() + " "
    return texto.strip()

# Dividir en frases
def dividir_en_frases(texto):
    frases = re.split(r'(?<=[.!?¡¿])\s+', texto)
    return [f.strip() for f in frases if f.strip()]

# Documento Word
doc_salida = Document()
doc_salida.add_heading("Limpieza y Clasificación de Emociones", level=1)

# Lista de resultados
datos_excel = []

# Procesar archivos
for archivo in os.listdir(ruta_entrada):
    if archivo.endswith(".docx"):
        ruta_completa = os.path.join(ruta_entrada, archivo)
        texto = leer_docx(ruta_completa)
        frases = dividir_en_frases(texto)

        doc_salida.add_heading(f"Archivo: {archivo}", level=2)

        for i, frase in enumerate(frases, 1):
            resultado = emotion_pipeline(frase[:512])[0][0]
            emocion = resultado['label']
            score = round(resultado['score'], 3)

            doc_salida.add_paragraph(f"{i}. {frase}\n   ➤ Emoción: {emocion} (confianza: {score})")
            datos_excel.append({
                "Archivo": archivo,
                "Frase": frase,
                "Emocion": emocion,
                "Confianza": score,
                "Indice": i
            })

# Guardar Word
doc_salida.save(archivo_docx)

# Guardar Excel con gráficos
df = pd.DataFrame(datos_excel)
with pd.ExcelWriter(archivo_excel, engine='xlsxwriter') as writer:
    df.to_excel(writer, sheet_name='Emociones', index=False)

    # Gráfico 2: Pastel
    resumen = df['Emocion'].value_counts().reset_index()
    resumen.columns = ['Emocion', 'Frecuencia']
    resumen.to_excel(writer, sheet_name='Resumen', index=False)
    wb = writer.book
    ws = writer.sheets['Resumen']

    pie = wb.add_chart({'type': 'pie'})
    pie.add_series({
        'name': 'Distribución de emociones',
        'categories': ['Resumen', 1, 0, len(resumen), 0],
        'values':     ['Resumen', 1, 1, len(resumen), 1]
    })
    pie.set_title({'name': 'Gráfico circular de emociones'})
    ws.insert_chart('E2', pie)

    # Gráfico 3: Línea temporal
    df_linea = df[['Indice', 'Emocion']].copy()
    df_linea['EmocionIndex'] = df_linea['Emocion'].astype('category').cat.codes
    df_linea.to_excel(writer, sheet_name='Evolucion', index=False)
    ws3 = writer.sheets['Evolucion']
    linea = wb.add_chart({'type': 'line'})
    linea.add_series({
        'name': 'Evolución emocional',
        'categories': ['Evolucion', 1, 0, len(df_linea), 0],
        'values':     ['Evolucion', 1, 2, len(df_linea), 2],
        'marker': {'type': 'circle'}
    })
    linea.set_title({'name': 'Emociones a lo largo del texto'})
    linea.set_x_axis({'name': 'Frase'})
    linea.set_y_axis({'name': 'Emoción (codificada)'})
    ws3.insert_chart('E2', linea)

    # Gráfico 4: Heatmap (tabla)
    heatmap_data = df.groupby(['Archivo', 'Emocion']).size().unstack(fill_value=0)
    heatmap_data.to_excel(writer, sheet_name='Heatmap')

    # Gráfico 5: Burbuja como scatter con tamaños
    burbujas = df.groupby('Emocion').agg({'Confianza': 'mean', 'Frase': 'count'}).reset_index()
    burbujas.columns = ['Emocion', 'ConfianzaMedia', 'Frecuencia']
    burbujas.to_excel(writer, sheet_name='Burbujas', index=False)
    ws5 = writer.sheets['Burbujas']

    bubble_chart = wb.add_chart({'type': 'scatter', 'subtype': 'bubble'})
    bubble_chart.add_series({
        'name': 'Emociones',
        'categories': ['Burbujas', 1, 0, len(burbujas), 0],         # Emoción
        'values':     ['Burbujas', 1, 1, len(burbujas), 1],         # ConfianzaMedia
        'bubble_sizes': ['Burbujas', 1, 2, len(burbujas), 2],       # Frecuencia
        'data_labels': {'value': True}
    })
    bubble_chart.set_title({'name': 'Gráfico de burbujas: Emoción'})
    bubble_chart.set_x_axis({'name': 'Emoción'})
    bubble_chart.set_y_axis({'name': 'Confianza promedio'})
    bubble_chart.set_legend({'position': 'bottom'})
    ws5.insert_chart('E2', bubble_chart)

print(f"Word generado: {archivo_docx}")
print(f"Excel con gráficos (2–5) generado: {archivo_excel}")